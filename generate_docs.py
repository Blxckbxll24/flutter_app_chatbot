#!/usr/bin/env python3
"""
Script para generar documentación Word de ChatBot RAG Flutter
Convierte la documentación Markdown a formato .docx con estilos profesionales
"""

import os
import sys
from datetime import datetime

# Verificar si python-docx está instalado
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    print("✅ python-docx encontrado")
except ImportError:
    print("❌ python-docx no encontrado. Instalando...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn

def create_chatbot_documentation():
    """Crear documento Word con la documentación completa"""
    
    # Crear nuevo documento
    doc = Document()
    
    # Configurar estilos
    setup_styles(doc)
    
    # Portada
    create_cover_page(doc)
    
    # Índice (placeholder)
    create_table_of_contents(doc)
    
    # Contenido principal
    create_main_content(doc)
    
    # Guardar documento
    doc_path = "ChatBot_RAG_Flutter_Arquitectura.docx"
    doc.save(doc_path)
    print(f"✅ Documento generado: {doc_path}")
    
    return doc_path

def setup_styles(doc):
    """Configurar estilos personalizados del documento"""
    
    # Estilo para títulos principales
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(21, 101, 192)  # Azul principal
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Estilo para subtítulos
    subtitle_style = doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(18)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(211, 47, 47)  # Rojo principal
    subtitle_style.paragraph_format.space_before = Pt(12)
    subtitle_style.paragraph_format.space_after = Pt(6)
    
    # Estilo para código
    code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    
def create_cover_page(doc):
    """Crear portada del documento"""
    
    # Título principal
    title = doc.add_paragraph('📱 CHATBOT RAG FLUTTER', style='CustomTitle')
    title.paragraph_format.space_after = Pt(24)
    
    # Subtítulo
    subtitle = doc.add_paragraph('Documentación Técnica de Arquitectura', style='CustomSubtitle')
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    
    # Información del proyecto
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('📊 Proyecto:', 'ChatBot RAG Flutter'),
        ('🏗️ Arquitectura:', 'Clean Architecture + Provider'),
        ('🎯 Tecnología:', 'Flutter 3.9+ / Dart'),
        ('📅 Fecha:', datetime.now().strftime('%d de %B de %Y')),
        ('🔧 Versión:', 'v1.0.0'),
        ('👨‍💻 Autor:', 'Equipo de Desarrollo')
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
    
    # Salto de página
    doc.add_page_break()

def create_table_of_contents(doc):
    """Crear tabla de contenidos"""
    
    doc.add_paragraph('📋 ÍNDICE', style='CustomTitle')
    
    toc_items = [
        '1. Resumen Ejecutivo',
        '2. Arquitectura General', 
        '3. Componentes del Sistema',
        '4. Flujo de Datos',
        '5. Estructura de Archivos',
        '6. Patrones de Diseño',
        '7. Gestión de Estado',
        '8. Configuración y Entornos',
        '9. Interfaz de Usuario',
        '10. Servicios y APIs',
        '11. Consideraciones Técnicas',
        '12. Dependencias',
        '13. Deployment',
        '14. Conclusiones'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_page_break()

def create_main_content(doc):
    """Crear contenido principal del documento"""
    
    # 1. Resumen Ejecutivo
    create_section(doc, '🎯 RESUMEN EJECUTIVO', 
                  'ChatBot RAG Flutter es una aplicación móvil desarrollada en Flutter que implementa un sistema de chat inteligente con capacidades RAG (Retrieval-Augmented Generation).')
    
    create_bullet_list(doc, [
        '✅ Chat en tiempo real con interfaz moderna',
        '✅ Soporte para temas claro/oscuro', 
        '✅ Persistencia de historial local',
        '✅ Simulación de respuestas streaming',
        '✅ Configuración multi-entorno (.env)',
        '✅ Arquitectura escalable y mantenible'
    ])
    
    # 2. Arquitectura General
    create_section(doc, '🏗️ ARQUITECTURA GENERAL',
                  'La aplicación sigue una arquitectura Clean Architecture con separación por capas:')
    
    create_architecture_diagram(doc)
    
    # 3. Componentes del Sistema
    create_section(doc, '🔧 COMPONENTES DEL SISTEMA', '')
    
    create_component_section(doc, 'Core Components', [
        ('ApiClient', 'Cliente HTTP para comunicación con backend'),
        ('AppEnvironment', 'Gestión de configuración por entornos'),
        ('Message', 'Modelo de datos para mensajes del chat')
    ])
    
    create_component_section(doc, 'Services Layer', [
        ('HistoryService', 'Persistencia local con SharedPreferences'),
        ('ThemeService', 'Gestión de temas claro/oscuro'),
        ('StreamingService', 'Simulación de respuestas en streaming')
    ])
    
    create_component_section(doc, 'Feature: Chat', [
        ('ChatController', 'Gestión de estado con Provider'),
        ('ChatRepository', 'Capa de abstracción de datos'),
        ('ChatScreen', 'Interfaz principal del chat')
    ])
    
    # 4. Flujo de Datos
    create_section(doc, '🔄 FLUJO DE DATOS', 
                  'Proceso completo desde que el usuario envía un mensaje hasta recibir la respuesta:')
    
    create_flow_diagram(doc)
    
    # 5. Estructura de Archivos
    create_section(doc, '📁 ESTRUCTURA DE ARCHIVOS', '')
    create_file_structure(doc)
    
    # 6. Gestión de Estado
    create_section(doc, '📊 GESTIÓN DE ESTADO',
                  'Implementación del patrón Provider + ChangeNotifier para estado reactivo:')
    
    create_code_block(doc, '''
// Estado centralizado
class ChatController extends ChangeNotifier {
  List<Message> _messages = [];
  bool _loading = false;
  
  List<Message> get messages => List.unmodifiable(_messages);
  bool get loading => _loading;
  
  Future<void> send(String text) async {
    _loading = true;
    notifyListeners(); // Notifica cambios
    
    // ... lógica de negocio ...
    
    _loading = false;
    notifyListeners(); // Notifica completitud
  }
}

// Consumo en UI
Widget build(BuildContext context) {
  final controller = context.watch<ChatController>();
  
  return ListView.builder(
    itemCount: controller.messages.length,
    itemBuilder: (context, index) {
      return MessageBubble(message: controller.messages[index]);
    },
  );
}
    ''')
    
    # 7. Configuración y Entornos
    create_section(doc, '🌍 CONFIGURACIÓN Y ENTORNOS',
                  'Sistema de configuración dinámico basado en archivos .env:')
    
    create_env_configuration(doc)
    
    # 8. Interfaz de Usuario
    create_section(doc, '🎨 INTERFAZ DE USUARIO', '')
    create_ui_components(doc)
    
    # 9. APIs y Comunicación
    create_section(doc, '🌐 SERVICIOS Y APIs',
                  'Estructura de comunicación con el backend:')
    
    create_api_documentation(doc)
    
    # 10. Consideraciones Técnicas
    create_section(doc, '🔍 CONSIDERACIONES TÉCNICAS', '')
    create_technical_considerations(doc)

def create_section(doc, title, description):
    """Crear una nueva sección con título y descripción"""
    doc.add_paragraph(title, style='CustomSubtitle')
    if description:
        doc.add_paragraph(description)
    doc.add_paragraph()

def create_bullet_list(doc, items):
    """Crear lista de viñetas"""
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)

def create_architecture_diagram(doc):
    """Crear diagrama de arquitectura textual"""
    
    layers = [
        '┌─────────────────────────────────────────┐',
        '│              PRESENTATION               │',
        '│         (UI Components & Screens)       │',
        '├─────────────────────────────────────────┤',
        '│               BUSINESS                  │',
        '│         (Controllers & Logic)           │', 
        '├─────────────────────────────────────────┤',
        '│                 DATA                    │',
        '│      (Repositories & Data Sources)      │',
        '├─────────────────────────────────────────┤',
        '│              EXTERNAL                   │',
        '│         (APIs & Local Storage)          │',
        '└─────────────────────────────────────────┘'
    ]
    
    for layer in layers:
        p = doc.add_paragraph(layer, style='CodeBlock')

def create_component_section(doc, section_title, components):
    """Crear sección de componentes"""
    doc.add_paragraph(f'📦 {section_title}', style='Heading 3')
    
    for name, description in components:
        p = doc.add_paragraph()
        run = p.add_run(f'{name}: ')
        run.bold = True
        run.font.color.rgb = RGBColor(21, 101, 192)
        p.add_run(description)

def create_flow_diagram(doc):
    """Crear diagrama de flujo textual"""
    
    flow_steps = [
        '1. Usuario escribe → InputBar',
        '2. InputBar → ChatController.send()',
        '3. ChatController → ChatRepository.ask()',
        '4. ChatRepository → ApiClient.sendMessage()',
        '5. ApiClient → HTTP Request → Backend',
        '6. Backend → Response → ApiClient', 
        '7. ApiClient → StreamingService (opcional)',
        '8. StreamingService → UI Update',
        '9. UI Update → HistoryService.save()'
    ]
    
    for step in flow_steps:
        p = doc.add_paragraph(step, style='CodeBlock')
        p.paragraph_format.left_indent = Inches(0.25)

def create_file_structure(doc):
    """Crear estructura de archivos"""
    
    structure = '''
lib/
├── main.dart                          # 🚀 Punto de entrada
├── core/                             # ⚙️ Funcionalidad base
│   ├── api_client.dart              # 🌐 Cliente HTTP
│   ├── models.dart                  # 📦 Modelos de datos
│   ├── theme.dart                   # 🎨 Temas y estilos
│   └── services/                    # 🛠️ Servicios compartidos
│       ├── app_environment.dart     # 🌍 Gestión de entornos
│       ├── history_service.dart     # 💾 Persistencia local
│       ├── streaming_service.dart   # ⚡ Simulación streaming
│       └── theme_service.dart       # 🌓 Gestión de temas
└── features/                        # 🎯 Features de la app
    └── chat/                        # 💬 Feature de chat
        ├── chat_controller.dart     # 🎛️ Controlador de estado
        ├── chat_repository.dart     # 📊 Capa de datos
        ├── chat_screen.dart         # 📱 Pantalla principal
        └── widgets/                 # 🧩 Componentes UI
            ├── input_bar.dart       # ⌨️ Barra de entrada
            ├── message_bubble.dart  # 💭 Burbuja de mensaje
            ├── streaming_indicator.dart # ⚡ Indicador streaming
            └── typing_indicator.dart # ⏳ Indicador de typing
    '''
    
    create_code_block(doc, structure)

def create_env_configuration(doc):
    """Crear documentación de configuración de entornos"""
    
    doc.add_paragraph('Archivo .env.dev (Desarrollo):', style='Heading 4')
    create_code_block(doc, '''
ENVIRONMENT=development
API_BASE_URL=http://localhost:8000/
DEBUG_MODE=true
API_TIMEOUT=30000
MAX_HISTORY_MESSAGES=100
ENABLE_LOGGING=true
    ''')
    
    doc.add_paragraph('Archivo .env.prod (Producción):', style='Heading 4') 
    create_code_block(doc, '''
ENVIRONMENT=production
API_BASE_URL=https://api.chatbot-rag.com/
DEBUG_MODE=false
API_TIMEOUT=15000
MAX_HISTORY_MESSAGES=50
ENABLE_LOGGING=false
    ''')

def create_ui_components(doc):
    """Crear documentación de componentes UI"""
    
    components = [
        ('ChatScreen', 'Pantalla principal con AppBar, ListView de mensajes e InputBar'),
        ('MessageBubble', 'Componente de burbuja con estilos por rol (user/bot/system)'),
        ('InputBar', 'Barra de entrada con TextField y botón de envío con gradiente'),
        ('TypingIndicator', 'Animación de puntos cuando el bot está escribiendo'),
        ('StreamingIndicator', 'Muestra texto en tiempo real con cursor parpadeante')
    ]
    
    for name, description in components:
        doc.add_paragraph(f'🧩 {name}', style='Heading 4')
        doc.add_paragraph(description)
        doc.add_paragraph()

def create_api_documentation(doc):
    """Crear documentación de API"""
    
    doc.add_paragraph('Request al Backend:', style='Heading 4')
    create_code_block(doc, '''
POST /chat
{
  "message": "¿Qué es RAG?",
  "history": [
    {
      "role": "user",
      "content": "Hola",
      "timestamp": "2025-01-23T10:00:00Z"
    }
  ]
}
    ''')
    
    doc.add_paragraph('Response del Backend:', style='Heading 4')
    create_code_block(doc, '''
{
  "answer": "RAG es una técnica de IA...",
  "timestamp": "2025-01-23T10:00:10Z",
  "status": "success"
}
    ''')

def create_technical_considerations(doc):
    """Crear consideraciones técnicas"""
    
    considerations = [
        ('Performance', [
            'ListView.builder para lazy loading de mensajes',
            'Provider.watch para reconstrucción selectiva de widgets',
            'const constructors para evitar reconstrucciones innecesarias',
            'Timeouts apropiados para evitar bloqueos de UI'
        ]),
        ('Seguridad', [
            'Validación y sanitización de input del usuario',
            'Manejo seguro de errores sin exposición de información sensible',
            'URLs HTTPS en producción',
            'Timeouts para prevención de ataques DoS'
        ]),
        ('Mantenibilidad', [
            'Separación clara de responsabilidades por capa',
            'Documentación en código y naming descriptivo',
            'Organización por features para escalabilidad',
            'Configuración externa mediante archivos .env'
        ])
    ]
    
    for category, items in considerations:
        doc.add_paragraph(f'🔍 {category}:', style='Heading 4')
        for item in items:
            p = doc.add_paragraph(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.25)
        doc.add_paragraph()

def create_code_block(doc, code_text):
    """Crear bloque de código formateado"""
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph(line, style='CodeBlock')

if __name__ == "__main__":
    print("🚀 Generando documentación de ChatBot RAG Flutter...")
    
    try:
        doc_path = create_chatbot_documentation()
        print(f"✅ ¡Documentación generada exitosamente!")
        print(f"📄 Archivo: {doc_path}")
        print(f"📍 Ubicación: {os.path.abspath(doc_path)}")
        
        # Intentar abrir el documento automáticamente
        if os.name == 'nt':  # Windows
            os.startfile(doc_path)
        elif os.name == 'posix':  # macOS/Linux
            os.system(f'open "{doc_path}"' if sys.platform == 'darwin' else f'xdg-open "{doc_path}"')
            
    except Exception as e:
        print(f"❌ Error generando documentación: {e}")
        sys.exit(1)
